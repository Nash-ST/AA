"""
Markdown to Word Document Converter
Converts Chinese construction industry markdown templates to formatted .docx files
"""
import argparse
import markdown
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re
import sys


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = cell._element.get_or_add_tcPr()
    shd = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color.replace('#', ''),
        qn('w:val'): 'clear'
    })
    shading_elm.append(shd)


def add_formatted_table(doc, table_data, header_color='1E3A5F'):
    """添加格式化的表格"""
    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, row_data in enumerate(table_data):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = str(cell_text)
            
            # 设置字体
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = '宋体'
                    r = run._element
                    r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 表头特殊格式化
            if i == 0:
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                set_cell_shading(cell, header_color)
    
    # 设置表格边框
    tbl = table._element
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = borders.makeelement(qn(f'w:{border_name}'), {
            qn('w:val'): 'single',
            qn('w:sz'): '4',
            qn('w:space'): '0',
            qn('w:color'): '000000'
        })
        borders.append(border)
    tblPr.append(borders)
    
    return table


def parse_markdown_table(md_text):
    """解析markdown表格"""
    lines = md_text.strip().split('\n')
    if len(lines) < 3:
        return None
    
    # 检查是否有表头分隔线
    if '|' not in lines[1].strip() or '-' not in lines[1].strip():
        return None
    
    header = [cell.strip() for cell in lines[0].split('|') if cell.strip()]
    rows = []
    
    for line in lines[2:]:
        if not line.strip().startswith('|'):
            break
        row = [cell.strip() for cell in line.split('|') if cell.strip()]
        if row:
            rows.append(row)
    
    if rows:
        return [header] + rows
    return None


def convert_md_to_docx(md_content, output_path):
    """将markdown内容转换为Word文档"""
    doc = Document()
    
    # 设置默认中文字体
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(10.5)
    
    lines = md_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # 标题
        if line.startswith('# '):
            heading = line[2:].strip()
            p = doc.add_heading(heading, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            heading = line[3:].strip()
            p = doc.add_heading(heading, level=2)
        elif line.startswith('### '):
            heading = line[4:].strip()
            p = doc.add_heading(heading, level=3)
        elif line.startswith('#### '):
            heading = line[5:].strip()
            p = doc.add_heading(heading, level=4)
        
        # 表格
        elif line.startswith('|') and '---' not in line:
            table_md = '\n'.join(lines[i:])
            table_data = parse_markdown_table(table_md)
            if table_data:
                add_formatted_table(doc, table_data)
                # 跳过表格行
                while i < len(lines) and (lines[i].strip().startswith('|') or '---' in lines[i]):
                    i += 1
                continue
        
        # 无序列表
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        
        # 有序列表
        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(re.sub(r'^\d+\.\s', '', line), style='List Number')
        
        # 空行
        elif not line:
            i += 1
            continue
        
        # 普通段落
        elif line:
            p = doc.add_paragraph(line)
        
        i += 1
    
    # 保存文档
    doc.save(output_path)
    print(f"Document saved to: {output_path}")


def main():
    parser = argparse.ArgumentParser(description='Convert Markdown to Word for Chinese construction documents')
    parser.add_argument('--input', '-i', required=True, help='Input markdown file path')
    parser.add_argument('--output', '-o', required=True, help='Output Word file path')
    
    args = parser.parse_args()
    
    with open(args.input, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    convert_md_to_docx(md_content, args.output)


if __name__ == '__main__':
    main()
